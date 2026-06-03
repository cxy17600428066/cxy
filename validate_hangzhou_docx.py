from pathlib import Path
import sys
import zipfile

from docx import Document

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

base = Path.cwd() / "杭州楼市26年物料"
ok = True

for p in sorted(base.glob("*.docx")):
    if not zipfile.is_zipfile(p):
        print("BAD_ZIP", p.name)
        ok = False
        continue
    try:
        doc = Document(p)
    except Exception as exc:
        print("BAD_DOCX", p.name, exc)
        ok = False
        continue
    text = "\n".join(x.text for x in doc.paragraphs)
    if not text.strip():
        print("EMPTY", p.name)
        ok = False
    if "联网校正" not in text and "联网更新说明" not in text:
        print("MISSING_UPDATE_MARK", p.name)
        ok = False
    print("OK", p.name, "paragraphs", len(doc.paragraphs), "tables", len(doc.tables))

policy = Document(base / "2026年杭州买房全政策详细分析.docx")
policy_text = "\n".join(p.text for p in policy.paragraphs)
required = [
    "不再审核购房资格",
    "最高额度由130万元提高到180万元",
    "5年期以上3.50%",
    "按3%征收率",
    "9087套",
]
for item in required:
    if item not in policy_text:
        print("MISSING_POLICY_TEXT", item)
        ok = False

raise SystemExit(0 if ok else 1)
