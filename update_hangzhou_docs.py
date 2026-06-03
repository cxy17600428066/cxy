from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE = Path.cwd() / "杭州楼市26年物料"
UPDATED_AT = "2026-06-03"


SOURCES = [
    ("杭州全面取消住房限购", "新华社/新华网，2024-05-09", "https://www.news.cn/20240509/b39dc5497d0e45ebbd9dbcbd762afba5/c.html"),
    ("杭州2024年10月楼市新政解读", "浙江在线转杭州市住保房管局解读，2024-10-10", "https://zjnews.zjol.com.cn/zjnews/202410/t20241010_30571193.shtml"),
    ("商业性个人住房贷款最低首付比例", "中国人民银行、国家金融监督管理总局，2024-09-24", "https://www.pbc.gov.cn/goutongjiaoliu/113456/113469/2025092212554216285/index.html"),
    ("杭州住房公积金使用政策", "杭州市人民政府公报，2026-03-30", "https://zfgb.hangzhou.gov.cn/11/103220263/t130220263034/530322.shtml"),
    ("个人购房契税政策", "国家税务总局热点问答，2024-12-13", "https://shanxi.chinatax.gov.cn/web/detail/sx-11400-548-1802372"),
    ("个人销售住房增值税政策", "财政部、税务总局公告2025年第17号，2025-12-29", "https://guizhou.chinatax.gov.cn/wjjb/zcfgk/szfl/zzs/202512/t20251231_89105624.html"),
    ("2026年5月LPR", "中新网据央行网站，2026-05-20", "https://www.chinanews.com.cn/cj/2026/05-20/10624697.shtml"),
    ("2026年5月杭州二手房成交", "第一财经，2026-06-02", "https://www.yicai.com/news/103211762.html"),
    ("2026年4月杭州二手房成交", "杭州网房产频道，2026-05-30", "https://house.hangzhou.com.cn/content/2026-05/30/content_9231466.html"),
]


def set_run_font(run, size=None, bold=None, color=None):
    font = run.font
    font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if color is not None:
        font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level <= 3 else "Normal"
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True, color=(31, 78, 121))
    return p


def add_body(doc, text="", bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_table(doc, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, val in enumerate(rows[0]):
        hdr[i].text = val
        for run in hdr[i].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color=(255, 255, 255))
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F4E79")
        tc_pr.append(shd)
    for row in rows[1:]:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
    doc.add_paragraph()
    return table


def add_sources(doc):
    add_heading(doc, "来源清单", 2)
    for title, org, url in SOURCES:
        add_bullet(doc, f"{title}：{org}。{url}")


def new_doc(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Pt(54)
    sec.bottom_margin = Pt(54)
    sec.left_margin = Pt(58)
    sec.right_margin = Pt(58)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, size=18, bold=True, color=(31, 78, 121))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, size=10, color=(89, 89, 89))
    doc.add_paragraph()
    return doc


def save_by_title(title_startswith, doc):
    for path in BASE.glob("*.docx"):
        try:
            old = Document(path)
            title = old.paragraphs[0].text.strip() if old.paragraphs else ""
        except Exception:
            continue
        if title.startswith(title_startswith):
            doc.save(path)
            return path
    raise FileNotFoundError(title_startswith)


def build_policy_doc():
    doc = new_doc(
        "2026年杭州买房政策联网校正版",
        f"更新日期：{UPDATED_AT}；已按公开来源修正限购、限售/摇号、贷款、公积金、税费口径",
    )
    add_body(doc, "本版替换原文中互相冲突或无法核实的2026年1月/2月政策表述。涉及实际购房资格、贷款、税费、补贴申请时，仍需以杭州市住保房管局、税务、公积金中心和经办银行实时审核为准。", bold=True)

    add_heading(doc, "一、限购与购房资格", 1)
    add_bullet(doc, "最新可核实口径：杭州已于2024年5月9日全面取消住房限购，在杭州市范围内购买住房不再审核购房资格。")
    add_bullet(doc, "不再区分上城、拱墅、西湖、滨江等中心城区和非中心城区的购房资格限制；原资料中“四区仍限购”的表述已删除。")
    add_bullet(doc, "购房落户：在杭州取得合法产权住房的非杭州户籍人员，可按政策申请落户；具体落户材料和流程以公安、政务服务平台实时要求为准。")

    add_heading(doc, "二、限售、限价与公证摇号", 1)
    add_bullet(doc, "2024年5月9日起，购房意向登记家庭数量小于或等于准售房源数量的新建商品住房项目，取消公证摇号销售要求，由开发企业自主销售。")
    add_bullet(doc, "2024年10月9日后出让的住宅用地，执行“新地新办法”：新建商品房不再设置限价；若报名家庭数大于准售房源数，仍需公证摇号，但不再实行社保排序和限售。")
    add_bullet(doc, "2024年10月9日前已出让或公告出让地块，仍按原合同、公告和项目销售规则执行。原资料中“热点盘5年、普通盘3年、法拍房5年”的统一表述不再作为最新通用口径。")

    add_heading(doc, "三、贷款与利率", 1)
    add_table(doc, [
        ["项目", "联网校正后的口径", "操作提示"],
        ["商业贷款首付", "全国层面商业性个人住房贷款不再区分首套、二套，最低首付款比例统一为不低于15%。", "银行会结合征信、收入、负债、房龄和房屋类型审批。"],
        ["商业贷款利率", "2026年5月20日LPR：1年期3.00%，5年期以上3.50%。", "实际房贷利率不是固定3.05%，以银行当日加点/减点为准。"],
        ["公积金贷款额度", "杭州公积金最高额度由130万元提高到180万元，个人最高90万元；个人可贷额度倍数由15倍调至20倍。", "2026年4月1日起施行，具体额度以公积金中心审核为准。"],
        ["公积金额度上浮", "新市民/青年家庭上浮20%，多子女和高层次人才家庭上浮50%，绿色低碳建筑或以旧换新上浮20%；符合多类条件可择高叠加，最高70%。", "是否可叠加、是否触及上限，以公积金中心细则为准。"],
    ])

    add_heading(doc, "四、交易税费", 1)
    add_table(doc, [
        ["税种", "联网校正后的口径", "备注"],
        ["契税", "家庭唯一住房：140㎡及以下1%，140㎡以上1.5%；家庭第二套住房：140㎡及以下1%，140㎡以上2%。", "2024年12月1日起执行；三套及以上、非住宅等按税务部门核定。"],
        ["增值税", "个人销售购买不足2年的住房，按3%征收率全额缴纳增值税；购买2年以上（含2年）对外销售，免征增值税。", "2026年1月1日起施行。"],
        ["个人所得税", "满五唯一通常可免征；不符合免征条件的，按地方税务核定方式办理。", "具体以杭州税务窗口核算为准。"],
        ["换房退税", "居民换购住房个人所得税退税政策需以财政部、税务总局延续文件和杭州税务办理口径核验。", "申请前先确认旧房、新房时间窗口和家庭住房套数。"],
    ])

    add_heading(doc, "五、买房决策提示", 1)
    add_bullet(doc, "政策门槛降低不等于资产价格确定上涨。2026年4-5月杭州二手房成交活跃，5月全市二手房网签9087套，连续三个月突破9000套，但市场仍呈现改善盘、核心板块和普通板块分化。")
    add_bullet(doc, "实际筛选房源时，应同时核验：挂牌量、近90天成交价、同小区去化速度、学区预警、地铁/产业兑现、物业和房龄。")
    add_sources(doc)
    return doc


def build_subsidy_doc():
    doc = new_doc(
        "2026年杭州市全域房产政策及补贴联网校正版",
        f"更新日期：{UPDATED_AT}；重点修正限购、限售、公积金、税费和市场判断",
    )
    add_body(doc, "本文件将原“全域政策及补贴大全”中已过时或互相冲突的基础政策改为联网可核验口径。区县补贴更新频率高，特别是消费券、限时补助、指定楼盘补贴，必须在网签前向项目属地住建局或开发企业书面核验。", bold=True)

    add_heading(doc, "一、全域基础政策", 1)
    add_table(doc, [
        ["主题", "最新校正"],
        ["限购", "杭州全市范围购房不再审核购房资格；原“四区保留限购”口径删除。"],
        ["新房摇号", "登记家庭数小于或等于房源数的项目取消公证摇号，由开发企业自主销售；超过房源数的仍按规则公证摇号。"],
        ["限售/限价", "2024年10月9日后出让住宅用地不再设置新房限价；报名人数大于准售房源数仍摇号，但不再社保排序和限售。此前地块按原合同/公告执行。"],
        ["商贷", "最低首付不低于15%；利率随LPR和银行加点/减点变化，不再写死为3.05%。"],
        ["公积金", "2026年4月1日起，杭州公积金最高贷款额度提高至180万元，个人最高90万元，可贷倍数20倍。"],
        ["税费", "契税按140㎡分档；2026年1月1日起个人销售未满2年住房增值税征收率为3%，满2年免征。"],
    ])

    add_heading(doc, "二、补贴信息处理原则", 1)
    add_bullet(doc, "市级人才、公租房、各区县购房消费券和房票安置政策仍保留为线索，但不再作为“确定可享受金额”直接引用。")
    add_bullet(doc, "桐庐、淳安、建德、钱塘、萧山、临平、富阳、临安等区县补贴多为限时、限额、指定项目或先到先得，申请前需核验有效期、楼盘清单、资金池余额、是否可叠加。")
    add_bullet(doc, "同一房产通常不能重复享受同类型补贴；已享受补贴后退房，一般需退回现金或注销/退回消费券。")
    add_bullet(doc, "若补贴要求“新建商品住宅”，商业、办公、公寓类非住宅通常不适用；二手房是否适用应逐条核验。")

    add_heading(doc, "三、2026年市场判断补充", 1)
    add_bullet(doc, "公开报道显示，杭州2026年4月市区二手房成交9968套，环比上涨6.5%、同比上涨5.8%；5月全市二手房网签9087套，同比大涨近18%。")
    add_bullet(doc, "成交回暖主要体现为结构性活跃：核心改善、优质学区/次新、总价适配刚需的产品更容易成交，远郊、老旧、无明确配套兑现的房源仍需重视流动性风险。")

    add_heading(doc, "四、已修正的原资料风险点", 1)
    add_bullet(doc, "删除“上城、拱墅、西湖、滨江四区仍限购”的口径。")
    add_bullet(doc, "删除“热点盘统一限售5年、普通盘统一限售3年、法拍房统一限售5年”的通用结论，改为按地块出让时间和项目规则核验。")
    add_bullet(doc, "将公积金最高额度从130万元更新为180万元。")
    add_bullet(doc, "将商业贷款利率从固定值改为“参考LPR+银行实际加减点”。")
    add_bullet(doc, "将个人销售未满2年住房增值税主税率更新为3%。")
    add_sources(doc)
    return doc


def build_resale_doc():
    doc = new_doc(
        "2026年杭州二手房交易流程联网校正版",
        f"更新日期：{UPDATED_AT}；按最新购房资格、贷款、公积金、税费口径修正",
    )
    add_heading(doc, "一、交易前准备", 1)
    for item in [
        "购房资格：杭州全市购房不再审核购房资格，但贷款、落户、税费和补贴仍需按个人情况审核。",
        "预算：总预算=成交价+契税+可能的增值税/个税议价转嫁+中介费+评估费+装修/维修成本。",
        "贷款预审：提前查征信、收入流水、负债和首付款来源。商贷最低首付不低于15%，但银行审批可能提高首付或调整利率。",
        "公积金：2026年4月1日起最高额度180万元、个人最高90万元，能否贷满取决于缴存余额、缴存年限、房屋情况和家庭贷款记录。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "二、看房与产权核验", 1)
    for item in [
        "核验不动产权证、共有权人、抵押、查封、租赁、户口占用、学位使用和物业欠费。",
        "优先使用浙里办、杭州房产交易相关平台或不动产窗口核验，避免只看中介截图。",
        "对学区房、老旧小区、顶底楼、临街房、回迁房、法拍/司法相关房源，单独做风险清单。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、签约、网签与资金监管", 1)
    for item in [
        "合同写清成交价、付款节点、贷款未批处理方式、户口迁出、交房时间、家具家电、违约责任。",
        "首付款、尾款优先进入资金监管账户，不建议直接转给房东或中介个人账户。",
        "若卖方房屋有未结清贷款，需明确解押资金来源、监管方式、解押时限和逾期责任。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "四、缴税过户与放款交房", 1)
    add_table(doc, [
        ["事项", "最新口径/操作"],
        ["契税", "按140㎡和家庭套数分档；家庭唯一/第二套140㎡及以下均为1%。"],
        ["增值税", "未满2年按3%征收率全额缴纳，满2年免征；实缴以税务核算为准。"],
        ["个税", "满五唯一通常免征；否则按税务核定方式办理，并在议价阶段确认由谁承担。"],
        ["放款交房", "银行放款后办理物业、水电燃气、维修基金、钥匙门禁、车位、家具家电清点。"],
    ])

    add_heading(doc, "五、2026年市场操作提示", 1)
    add_bullet(doc, "杭州二手房2026年5月网签9087套，市场活跃度高于近年同期，但成交分化明显；不要用全市成交量替代具体小区成交价。")
    add_bullet(doc, "买入前至少对比同小区近3-6个月真实成交、挂牌量、挂牌到成交周期和同板块新房竞争。")
    add_sources(doc)
    return doc


def insert_update_note(path, note):
    doc = Document(path)
    paras = doc.paragraphs
    # Avoid duplicate notes if rerun.
    for p in paras:
        if p.text.strip().startswith("【联网校正】"):
            p.text = note
            for run in p.runs:
                set_run_font(run, size=10, bold=True, color=(192, 0, 0))
            doc.save(path)
            return
    if not paras:
        p = doc.add_paragraph(note)
    else:
        p = paras[0].insert_paragraph_before(note)
        # Put note after title by moving the original title above it.
        title = paras[1].text if len(paras) > 1 else ""
    # If insert_paragraph_before was used, the note is before title; rebuild first two paragraphs.
    if len(doc.paragraphs) >= 2 and doc.paragraphs[0].text.startswith("【联网校正】"):
        title_text = doc.paragraphs[1].text
        doc.paragraphs[0].text = title_text
        doc.paragraphs[1].text = note
        for run in doc.paragraphs[0].runs:
            set_run_font(run, size=15, bold=True, color=(31, 78, 121))
        for run in doc.paragraphs[1].runs:
            set_run_font(run, size=10, bold=True, color=(192, 0, 0))
    doc.save(path)


def build_update_log():
    doc = new_doc("杭州楼市资料联网更新说明", f"更新日期：{UPDATED_AT}")
    add_heading(doc, "一、本次已修改文件", 1)
    add_bullet(doc, "重写：2026年杭州买房全政策详细分析.docx")
    add_bullet(doc, "重写：2026年杭州市全域房产政策及房产补贴大全.docx")
    add_bullet(doc, "重写：2026年杭州二手房交易流程1.docx")
    add_bullet(doc, "加注：学区、各区整体发展布局、杭州市发展规划等其余文档。")
    add_heading(doc, "二、核心修改", 1)
    for item in [
        "限购：统一修正为杭州全市购房不再审核购房资格。",
        "限售/限价/摇号：按2024年10月9日后出让住宅用地的新规则修正，不再使用原文热点盘/普通盘统一限售年限。",
        "公积金：最高额度更新为180万元，个人最高90万元，倍数20倍，并补充额度上浮规则。",
        "贷款：商业贷款最低首付保留15%，利率改为参考LPR和银行加减点。",
        "税费：契税按140㎡分档，个人销售未满2年住房增值税主税率更新为3%。",
        "市场：补充2026年4月9968套、5月9087套的杭州二手房成交信息，并提示结构性分化。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "三、未直接覆盖的内容", 1)
    add_body(doc, "学区划片、区县补贴、指定楼盘消费券、具体项目摇号/限售规则变化非常快，本次只做资料开头的联网校正提示，具体交易前仍需逐项核验官方公告。", bold=True)
    add_sources(doc)
    doc.save(BASE / "联网更新说明_20260603.docx")


def main():
    save_by_title("2026年杭州买房全政策详细分析", build_policy_doc())
    save_by_title("2026年杭州市全域房产政策及房产补贴大全", build_subsidy_doc())
    save_by_title("2026年杭州二手房交易流程", build_resale_doc())

    planning_note = (
        "【联网校正】截至2026-06-03，公开报道显示杭州二手房成交活跃度回升："
        "2026年4月市区成交9968套，5月全市网签9087套，连续三个月突破9000套。"
        "但市场呈结构性分化，本文件中的区县规划应作为产业/配套判断线索，不能直接等同于房价上涨结论；"
        "具体买房需叠加成交价、挂牌量、房龄、学区、地铁和项目规则核验。"
    )
    school_note = (
        "【联网校正】截至2026-06-03，本文件的学校名单和预警信息保留为选房线索；"
        "学区划片、红黄预警、落户年限和一表生规则以各区教育局当年招生公告为准。"
        "购买学区房前必须核验房户一致、学位占用、户口迁出和近年调剂情况。"
    )
    for path in BASE.glob("*.docx"):
        title = Document(path).paragraphs[0].text.strip()
        if title.startswith("2026年杭州各区县学区分类"):
            insert_update_note(path, school_note)
        elif "整体发展布局" in title or title.startswith("杭州市发展规划"):
            insert_update_note(path, planning_note)

    build_update_log()
    print("updated")


if __name__ == "__main__":
    main()
