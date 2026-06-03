from pathlib import Path
import sys

from docx import Document

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

base = Path.cwd() / "杭州楼市26年物料"
for p in sorted(base.glob("*.docx")):
    doc = Document(p)
    text = "\n".join(x.text.strip() for x in doc.paragraphs if x.text.strip())
    first = next((x.text.strip() for x in doc.paragraphs if x.text.strip()), "")
    print(p.name)
    print("  title:", first[:80])
    print("  has_note:", "联网校正" in text, "has_180:", "180万元" in text, "has_9087:", "9087套" in text)
